#!/usr/bin/env python3
"""
Script to convert Markdown technical writeup to Word document
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re

def convert_markdown_to_word(md_file, docx_file):
    """Convert Markdown file to Word document"""
    
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Split content by headers
    lines = md_content.split('\n')
    current_paragraph = None
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            if level == 1:
                # Main title
                heading = doc.add_heading(text, level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add some spacing
                doc.add_paragraph()
                
            elif level == 2:
                # Section headers
                heading = doc.add_heading(text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            elif level == 3:
                # Subsection headers
                heading = doc.add_heading(text, level=2)
                
            elif level == 4:
                # Sub-subsection headers
                heading = doc.add_heading(text, level=3)
                
        # Handle code blocks
        elif line.startswith('```'):
            if current_paragraph:
                current_paragraph = None
            continue
            
        # Handle regular paragraphs
        else:
            if line.startswith('- ') or line.startswith('* '):
                # List items
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                # Numbered list items
                p = doc.add_paragraph(line[3:], style='List Number')
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
                current_paragraph = p
    
    # Save the document
    doc.save(docx_file)
    print(f"✅ Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    convert_markdown_to_word("WhisperSOS_Technical_Writeup.md", "WhisperSOS_Technical_Writeup.docx") 